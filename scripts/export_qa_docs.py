# -*- coding: utf-8 -*-
"""
export_qa_docs.py — 一鍵重新生成小Q固定問答的審核文件（txt + docx）

用法（專案根目錄執行）：
  .venv\\Scripts\\python.exe scripts\\export_qa_docs.py

輸出：
  docs/AI機器人固定問答範本.docx   Word 版（微軟正黑體）
  （純文字 .txt 版依使用者需求已停產；要恢復見檔尾 __main__ 註解）

內容直接讀系統實裝模組（backend/services/canned_qa.py、coin_facts.py），
與線上版本永遠一致。**改了問答庫或知識庫後，跑一次本腳本即可更新審核文件。**
"""
import sys, io, os
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from backend.services.canned_qa import QA_TEMPLATES, KNOWLEDGE_INTENTS
from backend.services.coin_facts import COIN_FACTS
from backend.services import canned_qa
from backend.services.app_db import get_coins

INTENT_DESC = {
    "origin":      "起源：創辦人／年份／背景故事＋定位",
    "supply_coin": "供應量：上限／通膨／銷毀機制＋稀缺性白話說明",
    "holders":     "購買人數：誠實說明無精確統計＋替代指標（昨日成交量、採用概況）",
    "purpose":     "用途：用途特色＋風險註記",
    "what_is":     "是什麼／小檔案：定位＋起源＋用途＋供應＋共識＋採用＋現價立場＋風險",
}
VARS_DOC = [
    ("{幣名} {代號} {價格} {日期}", "幣種與最新收盤價", "資料庫 prices"),
    ("{立場} {分數}", "規則引擎判讀（偏多/中性/偏空、0~100）", "規則引擎（6 因子）"),
    ("{漲跌1日} {漲跌7日} {漲跌30日}", "近期漲跌 %", "資料庫 prices"),
    ("{RSI} {MA20} {MA60} {MA200} {布林上軌} {布林下軌}", "技術指標數值", "資料庫 indicators"),
    ("{趨勢說明} {動能說明} {量能說明} {位置說明} {技術快照}", "規則引擎白話文字", "ai_analyst"),
    ("{風險清單} {操作參考} {短線摘要}", "風險 / 建議 / 1h 視角", "規則引擎 + 1h indicators"),
    ("{本幣新聞數} {本幣看多} {本幣看空} {單幣情緒分} {市場情緒分}", "新聞情緒統計", "news.db 彙總"),
    ("{恐懼貪婪值} {恐懼貪婪標籤}", "全市場情緒", "fear_greed 表"),
    ("{幣數} {支援清單} {時線清單}", "啟用幣種與時線幣種（自動同步後台設定）", "app_config / DB"),
]
FIELD = [("positioning", "定位"), ("origin", "起源"), ("purpose", "用途"),
         ("supply", "供應"), ("consensus", "共識機制"), ("adoption", "採用概況"),
         ("risk", "風險註記")]


# ── TXT 版 ───────────────────────────────────────────────────────────────────
def export_txt(out_path: Path):
    L = []
    w = L.append
    BAR, BAR2 = "=" * 62, "-" * 62
    w(BAR); w("小Q AI 機器人 — 固定問答庫（實裝版）")
    w("由系統程式直接匯出，內容與線上版本完全一致"); w(BAR); w("")
    w("【回答架構】")
    w("1. 問題先比對本問答庫（取最長命中的觸發關鍵詞）")
    w("2. 命中 → 模板 + 即時數據組出固定答案；已設 GPT 金鑰則由 GPT 潤飾優化")
    w("   （數字/立場不得更動），失敗自動退回固定答案")
    w("3. 沒命中 → GPT 自由回答（被數據錨定）或規則引擎摘要，不硬答")
    w("")
    w("【幣種判定】免選幣：自動從問題偵測幣種（含錯字容忍）；未追蹤的幣誠實回覆沒資料。")
    w("")
    w("【模板變數】答案裡的 {…} 在回答當下自動帶入即時數據：")
    for v, meaning, src in VARS_DOC:
        w(f"  {v:<44s} {meaning}（{src}）")
    w("")
    w(BAR); w(f"第一部分：固定問答（共 {len(QA_TEMPLATES)} 條）"); w(BAR)
    cur = None
    for i, (cat, example, kws, template) in enumerate(QA_TEMPLATES, 1):
        if cat != cur:
            cur = cat
            n = sum(1 for c, *_ in QA_TEMPLATES if c == cat)
            w(""); w(f"◆◆◆ 分類：{cat}（{n} 條）◆◆◆")
        w(""); w(f"{i:02d}. 【{example}】")
        w(f"    觸發關鍵詞：{kws.replace('｜', '、')}")
        w("    答案模板：")
        for line in template.strip().split("\n"):
            w(f"      {line}")
        w("    核可：□ 通過　□ 修訂（意見：＿＿＿＿＿＿＿＿＿＿）")
    w(""); w(BAR)
    w(f"第二部分：幣種知識問答（{len(KNOWLEDGE_INTENTS)} 類意圖 × {len(COIN_FACTS)} 幣知識庫）")
    w(BAR); w("")
    w("以下問題由「幣種知識庫」（第三部分）組答案；查無知識檔的幣自動交給下一層。")
    for intent, kws in KNOWLEDGE_INTENTS:
        w(""); w(f"◆ {INTENT_DESC[intent]}")
        w(f"  觸發關鍵詞：{'、'.join(kws)}")
    w(""); w(BAR)
    w(f"第三部分：幣種知識庫（{len(COIN_FACTS)} 幣，公開既定事實）"); w(BAR)
    zh_map = {c["symbol"]: f'{c.get("zh","")} {c.get("ticker","")}' for c in get_coins()}
    for sym, facts in COIN_FACTS.items():
        w(""); w(BAR2); w(f"■ {zh_map.get(sym, sym)}"); w(BAR2)
        for key, label in FIELD:
            w(f"  【{label}】{facts[key]}")
        w("  核可：□ 通過　□ 修訂（意見：＿＿＿＿＿＿＿＿＿＿）")
    w(""); w(BAR); w("撰寫準則（新增/修訂時遵守）"); w(BAR)
    w("1. 不得預測價格、不得保證獲利")
    w("2. 涉及買賣的答案固定加「非投資建議」提醒")
    w("3. 觸發關鍵詞避免過短的常見字（防誤觸）")
    w("4. 知識類內容限公開既定事實")
    w(""); w("審核簽名：＿＿＿＿＿＿＿＿　　日期：＿＿＿＿＿＿＿＿")
    out_path.write_text("\n".join(L), encoding="utf-8-sig")
    print(f"寫出: {out_path} ({os.path.getsize(out_path)/1024:.0f} KB)")


# ── DOCX 版 ──────────────────────────────────────────────────────────────────
def export_docx(out_path: Path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    ZH_FONT = "微軟正黑體"

    def new_doc():
        doc = Document()
        st = doc.styles["Normal"]
        st.font.name = "Microsoft JhengHei"
        st.font.size = Pt(10.5)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
        for h, sz in (("Heading 1", 17), ("Heading 2", 14), ("Heading 3", 12)):
            try:
                s = doc.styles[h]
                s.font.name = "Microsoft JhengHei"
                s.font.size = Pt(sz)
                if s.element.rPr is None:
                    s.element.get_or_add_rPr()
                rf = (s.element.rPr.rFonts if s.element.rPr.rFonts is not None
                      else s.element.rPr._add_rFonts())
                rf.set(qn("w:eastAsia"), ZH_FONT)
                s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            except KeyError:
                pass
        return doc

    def add_rich(p, text):
        parts = str(text).replace("`", "").split("**")
        for i, seg in enumerate(parts):
            if not seg:
                continue
            r = p.add_run(seg)
            r.bold = (i % 2 == 1)
            r.font.name = "Microsoft JhengHei"
            rPr = r._element.get_or_add_rPr()
            rf = rPr.rFonts if rPr.rFonts is not None else rPr._add_rFonts()
            rf.set(qn("w:eastAsia"), ZH_FONT)

    def add_table(doc, rows, font_size=9.5):
        t = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
        t.style = "Table Grid"
        for ri, r in enumerate(rows):
            for ci, c in enumerate(r):
                cell = t.cell(ri, ci)
                cell.text = ""
                add_rich(cell.paragraphs[0], c)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(font_size)
                        if ri == 0:
                            run.font.bold = True
        doc.add_paragraph()
        return t

    doc = new_doc()
    doc.add_heading("AI 機器人固定問答庫（實裝版・審核追認用）", level=1)
    add_rich(doc.add_paragraph(),
             "本文件由系統程式直接生成，內容與線上實裝版完全一致。"
             "請逐條追認／標註修訂，修訂後由工程師同步進程式並重新生成本文件"
             "（指令：python scripts/export_qa_docs.py）。")
    doc.add_heading("一、回答架構：豐富固定答案為骨幹，GPT 只做優化", level=2)
    add_table(doc, [
        ["步驟", "行為", "特性"],
        ["① 固定問答庫比對", f"{len(QA_TEMPLATES)} 條模板 + {len(KNOWLEDGE_INTENTS)} 類幣種知識意圖；取最長命中關鍵詞", "內容可審核、即時數據填空、零幻覺"],
        ["② GPT 優化（若設金鑰）", "以固定答案為底稿潤飾", "數字/立場原封保留；失敗退回固定答案"],
        ["③ 沒命中", "GPT 自由回答（數據錨定）或規則引擎摘要", "不硬答"],
    ])
    doc.add_heading("二、答案模板的變數", level=2)
    add_table(doc, [["變數", "意義", "資料來源"]] + [list(v) for v in VARS_DOC], font_size=9)
    doc.add_heading(f"三、固定問答庫（{len(QA_TEMPLATES)} 條，請逐條追認）", level=2)
    rows = [["#", "分類", "問題示例", "觸發關鍵詞", "答案模板", "追認"]]
    for i, (cat, q, kw, ans) in enumerate(QA_TEMPLATES, 1):
        rows.append([str(i), cat, q, kw.replace("｜", "、"), ans, "□"])
    t = add_table(doc, rows, font_size=8.5)
    for w_, col in zip((0.8, 1.2, 3.0, 3.4, 8.0, 0.9), range(6)):
        for row in t.rows:
            row.cells[col].width = Cm(w_)
    doc.add_heading(f"四、幣種知識類意圖（{len(KNOWLEDGE_INTENTS)} 類 × {len(COIN_FACTS)} 幣）", level=2)
    add_table(doc, [["意圖", "觸發關鍵詞（節錄）", "回答內容"]] + [
        [intent, "、".join(kws[:8]) + "…", INTENT_DESC[intent].split("：", 1)[1]]
        for intent, kws in KNOWLEDGE_INTENTS], font_size=9)
    doc.add_heading("五、實際輸出示範（真實引擎渲染）", level=2)
    for q in ("現在適合進場嗎？", "為什麼跌？", "LTC 是什麼？"):
        r = canned_qa.try_answer("BTCUSDT" if "LTC" not in q else "LTCUSDT", q)
        add_rich(doc.add_paragraph(), f"❓ **{q}**")
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        add_rich(p, (r["answer"] if r else "（未命中）").replace("\n", "　⏎ "))
    doc.add_heading("六、新增／修訂申請表", level=2)
    rows = [["#", "問題", "觸發關鍵詞", "答案模板", "核可"]] + \
           [[str(i), "", "", "", "□"] for i in range(1, 9)]
    add_table(doc, rows)
    doc.add_heading("七、幣種知識庫全文", level=2)
    zh_map = {c["symbol"]: f'{c.get("zh","")} {c.get("ticker","")}' for c in get_coins()}
    for sym, facts in COIN_FACTS.items():
        doc.add_heading(zh_map.get(sym, sym), level=3)
        add_table(doc, [[label, facts[key]] for key, label in FIELD], font_size=9)
    add_rich(doc.add_paragraph(), "審核簽名：＿＿＿＿＿＿＿＿　　日期：＿＿＿＿＿＿＿＿")
    doc.save(str(out_path))
    print(f"寫出: {out_path} ({os.path.getsize(out_path)/1024:.0f} KB)")


if __name__ == "__main__":
    docs = ROOT / "docs"
    # 純文字 .txt 版依使用者需求不再產生（只留 Word 版）；要恢復取消下一行註解即可。
    # export_txt(docs / "AI機器人固定問答庫.txt")
    export_docx(docs / "AI機器人固定問答範本.docx")
    print("完成。修訂問答庫（canned_qa.py / coin_facts.py）後重跑本腳本即可同步審核文件（Word 版）。")
