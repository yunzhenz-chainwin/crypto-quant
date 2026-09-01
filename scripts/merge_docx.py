# -*- coding: utf-8 -*-
"""
merge_docx.py — 把 docs/ 下的全部交接 .docx 合併成「一份完整合集」給主管/接手一次看完。

用 docxcompose（保留每份原有的表格 / 標題 / 樣式），每份之間插入分頁，
最前面加一頁封面 + 目錄。

用法（專案根目錄）：
  .venv\\Scripts\\python.exe scripts\\merge_docx.py
輸出：
  docs/crypto-quant_文件合集.docx
本腳本會**先自動重生全部可生成的來源 docx**（md2docx.py 轉 7 份 .md：主管摘要、
  README、docs/archive 的 部署與運維／API規格／成果匯報／訊號增準計畫／研究預測評估；
  export_qa_docs.py 產問答範本、export_forecast_metrics_docx.py 產模型指標），
  即 9 章來源＋前言主管摘要全數重生，唯一例外「情緒詞庫範本.docx」(手工來源、無生成腳本)。
  → 一鍵重生合集：改 README / 各 .md / canned_qa.py 後，直接跑本腳本即可。
  這些 docx 是中繼產物（已 gitignore，不入版，合併後自刪）；只有合集與情緒詞庫範本.docx 進版控。
"""
import io
import os
import subprocess
import sys
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
OUT  = DOCS / "crypto-quant_文件合集.docx"

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docxcompose.composer import Composer

ZH = "微軟正黑體"

# 完整交接順序：總覽 → 維運/介面/資料/開發 → 訊號研究 → 預測研究 → AI/內容 → 匯報/規劃
SECTIONS = [
    ("crypto-quant_專案說明.docx",   "壹、專案說明（README）"),
    ("部署與運維.docx",              "貳、部署與運維・開發指南"),
    ("API規格.docx",                 "參、API 規格・資料庫說明"),
    ("成果匯報.docx",                "肆、成果匯報・訊號研究記錄"),
    ("訊號增準計畫.docx",            "伍、訊號增準計畫（規則式＋ML）"),
    ("研究預測評估.docx",            "陸、研究預測評估（成績單與校準）"),
    ("預測模型指標報告.docx",        "柒、預測模型指標與參考門檻"),
    ("AI機器人固定問答範本.docx",     "捌、AI 機器人固定問答庫"),
    ("情緒詞庫範本.docx",            "玖、情緒詞庫範本"),
]

# 前言：主管 2 分鐘摘要，排在封面/目錄之後、壹章之前（不編號）
PREFACE_DOCX = "主管摘要.docx"


def _ea(run, size=None, bold=False, color=None):
    run.font.name = "Microsoft JhengHei"
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rf = rPr.rFonts if rPr.rFonts is not None else rPr._add_rFonts()
    rf.set(qn("w:eastAsia"), ZH)


def _regen_sources():
    """用子行程重生全部可生成來源（md2docx ×7 + export_qa ×1 + metrics ×1，
    涵蓋 9 章來源＋前言主管摘要）。情緒詞庫範本.docx 無生成腳本（手工來源），不在此列，需自身存在。"""
    here = Path(__file__).resolve().parent
    for script in (
        "md2docx.py",
        "export_qa_docs.py",
        "export_forecast_metrics_docx.py",
    ):
        print(f"  重生來源：{script} …")
        subprocess.run([sys.executable, str(here / script)], check=True)


def main():
    _regen_sources()
    missing = [fn for fn in [PREFACE_DOCX] + [f for f, _ in SECTIONS] if not (DOCS / fn).exists()]
    if missing:
        raise SystemExit(f"缺少來源檔（情緒詞庫範本.docx 需手動存在）：{missing}")

    # 封面（空白 master + 中文字型）
    master = Document()
    st = master.styles["Normal"]
    st.font.name = "Microsoft JhengHei"
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), ZH)

    _ea(master.add_paragraph().add_run("crypto-quant 文件合集"), size=26, bold=True,
        color=RGBColor(0x1F, 0x3A, 0x5F))
    _ea(master.add_paragraph().add_run(f"加密貨幣量化分析平台 — 完整交接文件包（合併 {len(SECTIONS)} 份）"),
        size=12, color=RGBColor(0x66, 0x66, 0x66))
    master.add_paragraph()
    _ea(master.add_paragraph().add_run("內容"), size=14, bold=True)
    _ea(master.add_paragraph().add_run("前言　主管摘要（建議先看，2 分鐘看懂）"), size=11, bold=True)
    for i, (_, title) in enumerate(SECTIONS, 1):
        _ea(master.add_paragraph().add_run(f"{i}. {title}"), size=11)
    _ea(master.add_paragraph().add_run(
        "（各份保留原始格式；每份自第二頁起、以分頁分隔。要調整內容請改各來源檔後重跑 merge_docx.py）"),
        size=9, color=RGBColor(0x88, 0x88, 0x88))

    # 依序附加，每份前插分頁；主管摘要作為前言排在最前面
    composer = Composer(master)
    master.add_page_break()
    composer.append(Document(str(DOCS / PREFACE_DOCX)))
    for fn, _ in SECTIONS:
        master.add_page_break()
        composer.append(Document(str(DOCS / fn)))

    composer.save(str(OUT))
    print(f"寫出: {OUT}  ({os.path.getsize(OUT) / 1024:.0f} KB)")

    # 清掉可重生的中繼 docx，讓 docs/ 只留「合集 + 情緒詞庫範本.docx」
    for fn, _ in SECTIONS:
        if fn != "情緒詞庫範本.docx":
            try:
                (DOCS / fn).unlink()
            except FileNotFoundError:
                pass
    try:
        (DOCS / PREFACE_DOCX).unlink()
    except FileNotFoundError:
        pass
    print("  已清理中繼 docx（docs/ 只留 合集 + 情緒詞庫範本.docx）")


if __name__ == "__main__":
    main()
