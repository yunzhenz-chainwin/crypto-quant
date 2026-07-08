# -*- coding: utf-8 -*-
"""
export_readme_docx.py — 把 README.md 轉成 Word (.docx) 給主管看。

本機無 pandoc，改用 python-docx 自行轉換；支援：標題、表格、程式碼區塊、
引用、清單、**粗體**、`行內碼`、分隔線，以及 **```mermaid 流程圖/架構圖**。

Mermaid 圖會用本機 mermaid-cli（透過 npx 自動安裝）渲染成 PNG 嵌進 Word，
所以主管開 Word 直接看得到圖。若渲染失敗（無 node/npx 或離線），自動退回
「〔見線上版〕」文字註記，docx 仍能正常產生。

需求：node + npx（渲染 mermaid 用；首次會自動下載 @mermaid-js/mermaid-cli）。
用法（專案根目錄執行）：
  .venv\\Scripts\\python.exe scripts\\export_readme_docx.py
輸出：
  docs/crypto-quant_專案說明.docx
※ 改了 README.md 後重跑本腳本即可同步 Word 版。
"""
import io
import os
import re
import struct
import subprocess
import sys
import tempfile
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
ROOT = Path(__file__).resolve().parent.parent
SRC  = ROOT / "README.md"
OUT  = ROOT / "docs" / "crypto-quant_專案說明.docx"

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

ZH_FONT = "微軟正黑體"
MONO    = "Consolas"


def _set_ea(run, latin=None):
    if latin:
        run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rf = rPr.rFonts if rPr.rFonts is not None else rPr._add_rFonts()
    rf.set(qn("w:eastAsia"), ZH_FONT)


def _new_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Microsoft JhengHei"
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
    for h, sz in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        try:
            s = doc.styles[h]
            s.font.name = "Microsoft JhengHei"
            s.font.size = Pt(sz)
            if s.element.rPr is None:
                s.element.get_or_add_rPr()
            rf = (s.element.rPr.rFonts if s.element.rPr.rFonts is not None
                  else s.element.rPr._add_rFonts())
            rf.set(qn("w:eastAsia"), ZH_FONT)
            s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        except KeyError:
            pass
    return doc


def add_inline(p, text):
    for seg in re.split(r"(`[^`]*`)", text):
        if not seg:
            continue
        if len(seg) >= 2 and seg[0] == "`" and seg[-1] == "`":
            r = p.add_run(seg[1:-1])
            _set_ea(r, MONO)
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB0, 0x3A, 0x2E)
        else:
            for j, part in enumerate(seg.split("**")):
                if not part:
                    continue
                r = p.add_run(part)
                r.bold = (j % 2 == 1)
                _set_ea(r)


def _cells(line):
    parts = line.strip().split("|")
    if parts and parts[0].strip() == "":
        parts = parts[1:]
    if parts and parts[-1].strip() == "":
        parts = parts[:-1]
    return [c.strip() for c in parts]


def _is_sep(line):
    return bool(re.match(r"^\s*\|?\s*:?-{2,}.*$", line)) and set(line.strip()) <= set("|-: ")


def add_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(len(t.columns)):
            cell = t.cell(ri, ci)
            cell.text = ""
            add_inline(cell.paragraphs[0], row[ci] if ci < len(row) else "")
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.font.bold = True
    doc.add_paragraph()


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(4)
    for k, ln in enumerate(lines):
        if k:
            p.add_run().add_break()
        r = p.add_run(ln)
        _set_ea(r, MONO)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_diagram_note(doc):
    """Mermaid 渲染失敗時的退路：放乾淨的引導註記。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run("〔流程圖：請於 GitHub 線上版 README 檢視渲染後的圖表〕")
    _set_ea(r)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── Mermaid → PNG（本機 mermaid-cli via npx；失敗回 None）─────────────────────
_PCONF = None      # puppeteer no-sandbox 設定檔（一次性）
_MMDC_DEAD = False  # 一旦確認 npx/mmdc 不可用就不再重試


def _png_size(path):
    try:
        with open(path, "rb") as f:
            head = f.read(24)
        if head[:8] == b"\x89PNG\r\n\x1a\n":
            return struct.unpack(">II", head[16:24])
    except Exception:
        pass
    return None


def _render_mermaid(src, idx, tmp: Path):
    global _PCONF, _MMDC_DEAD
    if _MMDC_DEAD:
        return None
    mmd = tmp / f"d{idx}.mmd"
    png = tmp / f"d{idx}.png"
    mmd.write_text(src, encoding="utf-8")
    if _PCONF is None:
        _PCONF = tmp / "pconf.json"
        _PCONF.write_text('{"args":["--no-sandbox"]}', encoding="utf-8")
    cmd = (f'npx -y @mermaid-js/mermaid-cli -i "{mmd}" -o "{png}" '
           f'-b white -s 2 -p "{_PCONF}"')
    try:
        subprocess.run(cmd, shell=True, capture_output=True, timeout=180)
    except Exception:
        _MMDC_DEAD = True
        return None
    if not png.exists():
        if idx == 1:            # 第一張就失敗＝環境不支援，之後別再試
            _MMDC_DEAD = True
        return None
    return png


def add_diagram(doc, src, idx, tmp):
    """優先渲染成圖片嵌入；失敗才放文字註記。"""
    png = _render_mermaid(src, idx, tmp)
    if not png:
        add_diagram_note(doc)
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    sz = _png_size(png)
    if sz and sz[1] > sz[0] * 1.35:      # 偏高的圖：限高，避免超出頁面
        run.add_picture(str(png), height=Cm(17))
    else:                                # 一般/偏寬：限寬
        run.add_picture(str(png), width=Cm(14))
    return True


def main():
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = _new_doc()
    tmp = Path(tempfile.mkdtemp(prefix="mmd_"))
    mermaid_n = 0
    rendered = 0

    i, n = 0, len(md)
    in_code, code_buf, code_lang = False, [], ""
    try:
        while i < n:
            line = md[i]

            # 程式碼區塊（```mermaid → 渲染成圖）
            if line.strip().startswith("```"):
                if in_code:
                    if code_lang == "mermaid":
                        mermaid_n += 1
                        if add_diagram(doc, "\n".join(code_buf), mermaid_n, tmp):
                            rendered += 1
                    else:
                        add_code(doc, code_buf)
                    code_buf, in_code, code_lang = [], False, ""
                else:
                    in_code = True
                    code_lang = line.strip()[3:].strip().lower()
                i += 1
                continue
            if in_code:
                code_buf.append(line)
                i += 1
                continue

            # 表格
            if line.strip().startswith("|") and "|" in line.strip()[1:]:
                block = []
                while i < n and md[i].strip().startswith("|"):
                    if not _is_sep(md[i]):
                        block.append(_cells(md[i]))
                    i += 1
                if block:
                    add_table(doc, block)
                continue

            # 標題
            m = re.match(r"^(#{1,6})\s+(.*)$", line)
            if m:
                lvl = min(len(m.group(1)), 3)
                add_inline(doc.add_heading(level=lvl), m.group(2))
                i += 1
                continue

            # 分隔線
            if line.strip() == "---":
                doc.add_paragraph().add_run("─" * 30).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                i += 1
                continue

            # 引用
            if line.startswith(">"):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                p.add_run("▍ ").font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                add_inline(p, line.lstrip(">").strip())
                i += 1
                continue

            # 清單
            mb = re.match(r"^\s*[-*]\s+(.*)$", line)
            if mb:
                add_inline(doc.add_paragraph(style="List Bullet"), mb.group(1))
                i += 1
                continue

            # 空行
            if not line.strip():
                i += 1
                continue

            # 一般段落
            add_inline(doc.add_paragraph(), line)
            i += 1

        doc.save(str(OUT))
    finally:
        import shutil
        shutil.rmtree(tmp, ignore_errors=True)

    note = f"（{rendered}/{mermaid_n} 張 mermaid 圖已渲染嵌入）" if mermaid_n else ""
    if mermaid_n and rendered == 0:
        note += "  ⚠️ 渲染失敗，已退回文字註記（檢查 node/npx 是否可用）"
    print(f"寫出: {OUT}  ({os.path.getsize(OUT)/1024:.0f} KB) {note}")


if __name__ == "__main__":
    main()
