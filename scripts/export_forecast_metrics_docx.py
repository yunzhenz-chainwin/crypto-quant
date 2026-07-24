# -*- coding: utf-8 -*-
"""Generate the forecast metric chapter used by the combined Word handbook.

This report is intentionally authored as native Word content.  It keeps the
human-facing metric interpretation in ``crypto-quant_文件合集.docx`` without a
separate Markdown deliverable.  ``merge_docx.py`` creates this intermediate
document, appends it to the handbook, and removes the intermediate afterwards.
"""

from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.forecasting import (  # noqa: E402
    MIN_OBSERVATIONS,
    MIN_OUTCOMES,
    MIN_READY_CONFIDENCE,
    MODEL_VERSION,
)

OUT = ROOT / "docs" / "預測模型指標報告.docx"
REPORT_DATE = "2026-07-24"
ZH_FONT = "微軟正黑體"
NAVY = "1F3A5F"
BLUE = "DCE6F1"
LIGHT_BLUE = "EEF4F8"
LIGHT_RED = "FCE8E6"
LIGHT_YELLOW = "FFF4CE"
LIGHT_GREEN = "E6F4EA"
WHITE = "FFFFFF"


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _font(run, *, size=9.0, bold=False, color=None, italic=False) -> None:
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts if r_pr.rFonts is not None else r_pr._add_rFonts()
    r_fonts.set(qn("w:eastAsia"), ZH_FONT)


def _set_paragraph_text(paragraph, text: str, *, size=9.0, bold=False, color=None) -> None:
    paragraph.clear()
    run = paragraph.add_run(str(text))
    _font(run, size=size, bold=bold, color=color)
    paragraph.paragraph_format.space_after = Pt(0)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_table_grid_widths(table, widths) -> None:
    """Fix both Word's table grid and cell widths to the requested cm values.

    Setting ``cell.width`` alone only updates ``w:tcW``.  Word and
    docxcompose can still lay out the table from the stale ``w:tblGrid``, so
    every grid column and the aggregate ``w:tblW`` must be fixed as well.
    """
    table.autofit = False
    twips = [int(Cm(width).twips) for width in widths]
    grid_columns = list(table._tbl.tblGrid.gridCol_lst)
    if len(grid_columns) != len(twips):
        raise ValueError("table grid column count does not match requested widths")
    for grid_column, width_twips in zip(grid_columns, twips):
        grid_column.set(qn("w:w"), str(width_twips))

    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(sum(twips)))


def _add_table(doc, headers, rows, widths, *, font_size=8.2, status_column=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_grid_widths(table, widths)

    header = table.rows[0]
    _set_repeat_table_header(header)
    for index, (cell, title) in enumerate(zip(header.cells, headers)):
        cell.width = Cm(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade(cell, NAVY)
        _set_cell_margins(cell)
        _set_paragraph_text(cell.paragraphs[0], title, size=font_size, bold=True, color=WHITE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    status_colors = {
        "不合格": LIGHT_RED,
        "接近隨機": LIGHT_RED,
        "無 lift": LIGHT_RED,
        "低於基準": LIGHT_RED,
        "不可評估": LIGHT_RED,
        "N/A": LIGHT_YELLOW,
        "僅 coverage 落在規劃帶": LIGHT_YELLOW,
    }
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for col_index, (cell, value) in enumerate(zip(cells, values)):
            cell.width = Cm(widths[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            _set_paragraph_text(cell.paragraphs[0], value, size=font_size)
            if row_index % 2:
                _shade(cell, "F7F9FB")
        if status_column is not None:
            status = str(values[status_column])
            if status in status_colors:
                _shade(cells[status_column], status_colors[status])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def _add_bullet(doc, text: str, *, level=0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    run = paragraph.add_run(text)
    _font(run, size=9.5)
    paragraph.paragraph_format.space_after = Pt(3)


def _add_note(doc, title: str, text: str, *, fill=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_grid_widths(table, (15.0,))
    cell = table.cell(0, 0)
    cell.width = Cm(15.0)
    _shade(cell, fill)
    _set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    title_run = paragraph.add_run(f"{title}：")
    _font(title_run, size=10, bold=True, color=NAVY)
    body_run = paragraph.add_run(text)
    _font(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _configure_styles(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal.font.size = Pt(9.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
    for style_name, size in (("Title", 24), ("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)


def build_document() -> Document:
    doc = Document()
    _configure_styles(doc)
    section = doc.sections[0]
    # docxcompose uses the master document's section properties.  Keep this
    # chapter portrait and fit every table inside the handbook's printable
    # width so it remains readable after composition.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    properties = doc.core_properties
    properties.title = "crypto-quant 預測模型指標與參考門檻"
    properties.subject = f"{MODEL_VERSION} point-in-time replay、動態可信度與內部研究門檻"
    properties.author = "crypto-quant"

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(title.add_run("預測模型指標與參考門檻"), size=24, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(
        subtitle.add_run(
            f"{MODEL_VERSION}｜point-in-time replay｜報告日期 {REPORT_DATE}"
        ),
        size=10.5,
        color="666666",
    )

    _add_note(
        doc,
        "目前結論",
        "現有模型尚無可依賴的方向預測優勢，不應部署或對外宣稱可預測漲跌。"
        "Brier、log loss、AP 與 majority accuracy 均未勝過正確基準；AUC、Balanced accuracy、MCC 接近隨機。"
        "Platt／Beta 校準結果維持 keep_identity。",
        fill=LIGHT_RED,
    )
    _add_note(
        doc,
        "門檻定位",
        "本章的固定數字是 crypto-quant 的 P1 內部研究里程碑，不是跨市場通用及格線、投資建議或獲利保證。"
        "正式判定必須逐一針對 model_version × horizon，以 exact-vintage、未觸碰樣本及 paired confidence interval 驗證。",
    )

    doc.add_heading("一、公開首頁與主題分類選幣", level=1)
    entry_rows = [
        ("公開入口", "所有非 /admin 路徑進入公開 App；不再先顯示 MarketOverview，也沒有『市場總覽』導覽入口。"),
        ("第一個畫面", "直接顯示 BTCUSDT（比特幣 BTC）詳細頁；即使幣種 API 尚未回傳，選幣器仍保留 BTCUSDT 作為穩定 fallback。"),
        ("主題分類", "沿用 CATEGORIES：全部、主流幣、公鏈平台、DeFi 基礎、支付轉帳、迷因幣；分類籤直接放在詳細頁選幣區。"),
        ("切換規則", "分類籤只篩選下拉選單；若目前幣不屬於新分類，自動切到該分類第一個可用幣。切回『全部』時保留目前幣。"),
        ("責任邊界", "首頁分類與幣種選擇是 UI 導覽／查詢狀態，不訓練、不校準、不寫入模型參數，也不改變歷史成績。"),
    ]
    _add_table(
        doc,
        ("項目", "穩定行為"),
        entry_rows,
        (3.0, 12.0),
        font_size=8.2,
    )

    doc.add_heading("二、後台模型成績的動態可信度判讀", level=1)
    _add_note(
        doc,
        "契約仍可演進，原則不可退讓",
        "後端 assessment 欄位與前端 fallback 呈現仍可逐步整合，但同一組 fail-closed 規則必須固定："
        "沒有成熟資料就顯示不可驗證；樣本不足只可描述；缺值顯示『—／樣本不足』，不得補 0、補綠燈或沿用舊 snapshot。",
        fill=LIGHT_YELLOW,
    )
    trust_rows = [
        ("資料來源", "每次查詢從不可變 forecast ledger 的 first-issued snapshot 與已解析 outcome 動態計算；不得用畫面快取或 frozen Word 數字代替目前資料。"),
        ("資料成熟度", "只計 target 已到期、outcome 已封存且 probability／baseline／日期可評分的 records；未成熟與 unresolved 只顯示數量，不進 Accuracy、Brier 等分母。"),
        ("最低評估證據", "正式 scorecard gate 為每個 model × horizon 至少 1,000 筆、180 個 issue dates；低於門檻為 insufficient_evidence。這和模型自身 120 根日線／30 筆成熟 outcome 的 ready gate 是兩件事。"),
        ("可信度判定", "依 data maturity、provenance、正確 forecast-time baseline、Brier skill 與 paired CI 共同動態判讀；不是固定百分比，也不是單一 Accuracy。"),
        ("判定階梯", "unverifiable → insufficient_evidence／diagnostic_only → not_better_than_baseline → promising_not_confirmed → release_review_eligible；只有全部 hard gates 通過才可進人工發布審查。"),
        ("篩選後呈現", "window、symbol 或非單一 horizon 的切片可以看診斷，但 promotion gate 必須標示 not_applicable；不可把最好看的切片包裝成整體可信度。"),
        ("自動刷新", "前端可週期重抓與顯示 data_as_of／generated_at；重新整理只重算成績，不得在背景修改模型或門檻。"),
    ]
    _add_table(
        doc,
        ("層次", "動態判讀原則"),
        trust_rows,
        (3.0, 12.0),
        font_size=7.8,
    )

    doc.add_heading("三、評估範圍與證據限制", level=1)
    _add_bullet(doc, "資料範圍：15 個幣種、1／5／10 日 horizon，合計 73,881 筆已成熟 replay observations。")
    _add_bullet(doc, "Issue dates：2021-10-30 至 2026-07-19；positive／non-up support 為 34,956／38,925。")
    _add_bullet(doc, "Forecast-time baseline 只使用預測發出當下已成熟的歷史結果，不能使用最終全樣本上漲率。")
    _add_bullet(doc, "vintage_exact=false：使用目前保存的 CSV 重播，不是每個歷史發出時間點的不可變原始快照。")
    _add_bullet(doc, "Ready observations=0，且 live ledger 尚無已成熟 ready outcomes；所有 ready-only 指標必須保持 null。")
    _add_bullet(doc, "因此現有結果只屬研究診斷，不構成 production、交易或獲利證據。")

    doc.add_heading("四、目前指標怎麼樣", level=1)
    doc.add_heading("4.1 指標定義與後台顯示規則", level=2)
    metric_definition_rows = [
        ("Accuracy", "threshold 0.5 下預測正確的比例", "必須和同切片 majority baseline 一起看；類別不平衡時不可單獨判斷。"),
        ("Precision", "預測 up 中實際 up 的比例 TP/(TP+FP)", "沒有 predicted-up 時分母為 0，顯示 null／樣本不足，不得填 0。"),
        ("Recall", "實際 up 中被抓到的比例 TP/(TP+FN)", "高 Recall 可能來自幾乎全部預測 up，必須連同 Specificity 與 coverage。"),
        ("F1", "Precision 與 Recall 的 harmonic mean", "不含 TN；必須連同 Balanced accuracy、MCC 與固定 threshold 閱讀。"),
        ("ROC-AUC", "正例排序高於負例的機率型 ranking 指標", "0.5 近似隨機；只有單一 outcome class 時為 null。"),
        ("Average Precision（AP）", "tie-preserving step-wise precision-recall area", "必須和同切片 positive prevalence 比；不可用 trapezoid PR-AUC 混稱。"),
        ("Brier score", "mean((p-y)^2)，同時評估校準與辨別", "越低越好；promotion 優先和相同 forecast IDs 的 point-in-time baseline paired 比較。"),
        ("Brier Skill Score", "1 - Brier_model/Brier_baseline", "大於 0 才表示勝過 baseline；仍須 baseline-model loss CI 下界大於 0。"),
        ("Log loss", "-mean(y log p + (1-y) log(1-p))", "越低越好，會重罰過度自信的錯誤；必須 paired 比較同一批 outcomes。"),
        ("ECE", "固定 10 個 equal-width bins 的 probability/observed-rate gap 加權平均", "越低通常越好，但對分箱敏感；不能在 Brier 或 AUC 惡化時單獨升級。"),
        ("SHAP", "對固定 feature model f(X) 的 attribution", "historical statistical baseline 沒有固定 feature schema/model object，故 N/A；它不是準確率。"),
    ]
    _add_table(
        doc,
        ("指標", "回答什麼", "正確解讀／null 規則"),
        metric_definition_rows,
        (3.2, 4.4, 7.4),
        font_size=7.1,
    )

    doc.add_heading("4.2 固定 replay 的目前數值", level=2)
    current_rows = [
        ("Proper", "Brier / baseline", "0.252671 / 0.251558", "模型較差；BSS = -0.4424%", "不合格"),
        ("Proper", "Log loss / baseline", "0.698954 / 0.696363", "模型較差", "不合格"),
        ("Calibration", "ECE（10 bins）", "0.035678", "約 3.57 個百分點的 bin-weighted gap；不可抵銷負 BSS", "不合格"),
        ("Threshold", "Accuracy / majority", "0.510551 / 0.526861", "低於永遠預測 non-up", "低於基準"),
        ("Threshold", "Precision", "0.477891", "僅比 prevalence 0.473139 高 0.48pp", "不合格"),
        ("Threshold", "Recall / Specificity", "0.372554 / 0.634477", "漏掉約 62.7% 的 up events", "不合格"),
        ("Threshold", "F1", "0.418699", "固定 threshold 0.5 的診斷值；不能單看", "不合格"),
        ("Threshold", "Balanced accuracy", "0.503515", "接近 0.5 隨機基準", "接近隨機"),
        ("Threshold", "MCC", "0.007275", "接近 0，幾乎沒有整體分類關聯", "接近隨機"),
        ("Ranking", "ROC-AUC", "0.504585", "接近 0.5 隨機排序", "接近隨機"),
        ("Ranking", "AP / prevalence", "0.470341 / 0.473139", "沒有 precision-recall lift", "無 lift"),
        ("Uncertainty", "80% interval", "coverage 82.6775%; width 22.9604pp; WIS 4.504830", "coverage 落在暫定帶，但尚無 interval baseline／cluster CI", "僅 coverage 落在規劃帶"),
        ("Evidence", "Paired Brier advantage CI", "[-0.003007, 0.000843]", "baseline loss − model loss 的 95% CI 跨 0", "不合格"),
        ("Explainability", "SHAP", "N/A", "目前不是具有固定 feature schema 的 ML model", "N/A"),
    ]
    _add_table(
        doc,
        ("類別", "指標", "目前數值", "正確解讀", "判定"),
        current_rows,
        (1.4, 2.0, 2.7, 7.2, 1.7),
        font_size=6.9,
        status_column=4,
    )

    doc.add_heading("五、各 horizon 是否一致", level=1)
    horizon_rows = [
        ("1 日", "24,692", "0.250876", "0.250380", "-0.1979%", "0.500261", "0.488347", "0.022802"),
        ("5 日", "24,632", "0.252440", "0.251355", "-0.4319%", "0.503384", "0.469506", "0.033197"),
        ("10 日", "24,557", "0.254709", "0.252947", "-0.6963%", "0.496079", "0.449828", "0.051114"),
    ]
    _add_table(
        doc,
        ("Horizon", "N", "Brier", "Baseline Brier", "BSS", "ROC-AUC", "AP", "ECE"),
        horizon_rows,
        (1.5, 1.4, 1.8, 2.2, 1.5, 2.2, 2.2, 2.2),
        font_size=7.0,
    )
    _add_note(
        doc,
        "Horizon 判讀",
        "三個 horizon 的 BSS 全為負；10 日 AUC 低於 0.5，且 Brier、log loss、ECE 最差。"
        "不能用單一 horizon 的 accuracy 或 specificity 掩蓋其他核心指標。",
        fill=LIGHT_YELLOW,
    )

    doc.add_heading("六、預計多少才有參考意義", level=1)
    target_rows = [
        ("證據量與 provenance", "pooled N=73,881；vintage_exact=false；ready N=0", "每個 model+horizon 至少 1,000 筆、180 個 issue dates、正負類各 100；完整 universe、唯一 forecast ID、exact immutable vintage", "前述條件設為 hard gate；另累積同規模 live/shadow matured outcomes"),
        ("Brier / BSS", "0.252671 / -0.4424%", "BSS > 0，且 issue-date block-bootstrap 的 baseline − model 95% CI 下界 > 0", "每個 horizon BSS ≥ +1%。1／5／10 日 Brier 約需 ≤0.247876／0.248841／0.250418"),
        ("Log loss", "0.698954；baseline 0.696363", "低於相同 forecast IDs 的 baseline，且 paired improvement 95% CI 下界 > 0", "先以 ≥0.5% relative improvement；依目前 pooled baseline 約 ≤0.692881"),
        ("ECE", "0.035678", "相同 paired OOS rows 上，equal-width 與 tie-preserving equal-mass 都優於 identity，且 Brier/AUC 不退步", "兩種 ECE 都 ≤0.02；≤0.01 是較強校準，但不可單獨升級"),
        ("ROC-AUC", "0.504585", "issue-date clustered block-bootstrap 95% CI 下界 >0.5，且各 horizon／regime 穩定", "point AUC ≥0.53 作第一個弱訊號里程碑；≥0.55 作較強目標"),
        ("Average Precision", "0.470341；prevalence 0.473139", "AP > 同切片 prevalence，且重抽時同步重算 prevalence 的 AP lift 95% CI 下界 >0", "AP 至少比 prevalence 高 5% relative；目前約 ≥0.496796"),
        ("Accuracy", "0.510551；majority 0.526861", "高於同切片 majority baseline，且 Recall／Specificity 不崩潰", ">0.526861，並同時通過 Balanced accuracy 與 MCC；不作單一 promotion gate"),
        ("Balanced accuracy", "0.503515", "95% CI 下界 >0.5", "point ≥0.52；較強目標 ≥0.55"),
        ("MCC", "0.007275", "95% CI 下界 >0", "point ≥0.05；較強目標 ≥0.10"),
        ("Precision / Recall / F1", "0.477891 / 0.372554 / 0.418699", "Threshold 只能在 training/validation 事先固定，並與相同 coverage／成本政策 baseline 比較", "暫定聯合篩選：Precision ≥0.493139、Recall ≥0.50、Specificity ≥0.50、F1 ≥0.50；仍須通過 BSS／BA／MCC／成本條件"),
        ("Ready coverage", "0%", "預先登錄 coverage floor，再評分至少 1,000 筆／180 dates 的 matured ready outcomes", "產品尚未定義頻率時，先用 ≥10% 作 shadow 可用性里程碑；不得看完測試再調低"),
        ("80% prediction interval", "coverage 82.6775%；width 22.9604pp；WIS 4.504830", "Coverage 落在預先登錄的 77%–83% 產品容許帶，且 WIS／width 優於同口徑 baseline", "保持 coverage 在規劃帶，同時讓 WIS relative improvement ≥5%；width 需分 horizon"),
        ("SHAP", "N/A", "先有固定 feature schema、model object、background data 與 frozen OOS set", "報告方向、mean |SHAP| 與跨 chronological folds 穩定性；沒有越高越好的門檻"),
    ]
    _add_table(
        doc,
        ("指標／條件", "目前", "最低有參考意義的證據", "下一階段內部目標"),
        target_rows,
        (2.3, 2.8, 5.1, 4.8),
        font_size=6.6,
    )
    _add_note(
        doc,
        "最重要的判定規則",
        "達到 point target 但 confidence interval 仍跨過基準，只能稱為『值得繼續研究』。"
        "最低可依賴證據仍是：優於正確 forecast-time baseline，且 paired CI 排除無改善。",
        fill=LIGHT_GREEN,
    )

    doc.add_heading("七、F1、Recall 與 SHAP 為什麼不能單獨判斷", level=1)
    _add_bullet(doc, "依目前 prevalence，永遠預測 up 可得到 Precision 0.473139、Recall 1.0、F1 0.642355，卻沒有任何辨別何時上漲的能力。")
    _add_bullet(doc, "因此 F1 或 Recall 的漂亮數字可能只是退化成單一類別策略；必須連同 Specificity、Balanced accuracy、MCC、coverage 與成本閱讀。")
    _add_bullet(doc, "SHAP 是模型 attribution，不是效能分數。它回答『模型為什麼這樣預測』，不回答『模型有多準』，所以沒有越高越好的 target。")
    _add_bullet(doc, "目前 historical-baseline-v2 沒有固定 X → f(X) feature schema 與 model object，SHAP 必須誠實標示 N/A。")

    doc.add_heading("八、真實控制參數、查詢篩選與校準 challenger", level=1)
    _add_note(
        doc,
        "參數名稱必須對得上程式",
        "下列值是目前程式實際存在的 model/evaluation controls；不代表都已開放後台 UI 即時修改。"
        "若未來提供調參介面，必須建立版本化設定與獨立 validation，不能直接覆寫 frozen scorecard 的測試條件。",
    )
    control_rows = [
        ("MIN_READY_CONFIDENCE", str(MIN_READY_CONFIDENCE), "模型 confidence score 的研究發布門檻（0–100 分），不是 p(up) threshold；低於此值就 abstain。"),
        ("MIN_OBSERVATIONS", str(MIN_OBSERVATIONS), "產生預測前至少需要的 completed daily bars；屬模型資料準備門檻，不是 scorecard promotion N。"),
        ("MIN_OUTCOMES", str(MIN_OUTCOMES), "同 regime 可用成熟 outcomes 的最低數；不足時 fallback／abstain，不得引用未成熟未來結果。"),
        ("Regime", "MA60 + trailing 20d return", "以 close 對 MA60 與過去 20 日報酬定義 bull／bear／sideways；改 60 或 20 就是模型版本變更。"),
        ("Classification threshold", "0.5", "只用來把 p(up) 轉成 up/non-up 以計算 Accuracy、Precision、Recall、F1；p=0.5 歸 up。不是 ready confidence 40。"),
        ("Scorecard evidence gate", "1,000 observations / 180 issue dates", "評估可信度門檻；和模型 120 bars／30 outcomes 分開，不得混為同一參數。"),
    ]
    _add_table(
        doc,
        ("控制項", "目前值", "作用與邊界"),
        control_rows,
        (3.6, 3.2, 8.2),
        font_size=7.5,
    )

    filter_rows = [
        ("window", "只限制 scorecard 的 target date 查詢窗；不改模型、不重估參數。", "window-filtered diagnostic slice；formal promotion gate 為 not_applicable。"),
        ("horizon", "選擇 1／5／10 日既有 forecast group；不是把模型的 horizon 調成新值。", "正式審查必須明確選一個 horizon，不能 pooled 後宣稱通過。"),
        ("symbol", "只選單一幣種查看診斷；不改跨幣資料或模型係數。", "symbol-filtered diagnostic slice；不可取最好幣種代表完整 universe。"),
        ("model_version", "只選既有 immutable model version 的成績。", "正式審查必須明確指定一個版本；新參數應建立新版本而非覆寫舊版。"),
    ]
    _add_table(
        doc,
        ("畫面／API 篩選", "實際作用", "Promotion 限制"),
        filter_rows,
        (3.0, 6.0, 6.0),
        font_size=7.5,
    )
    _add_note(
        doc,
        "禁止 scorecard feedback loop",
        "Frozen scorecard 是獨立稽核／test 證據，不能由系統看完成績後自動改 MIN_READY_CONFIDENCE、MA window、classification threshold、"
        "calibrator 或任何模型參數再重跑同一份資料。調參只能在 training／validation 或 nested chronological walk-forward 內進行；"
        "參數鎖定後才可一次評估 untouched test，並建立新 model_version。",
        fill=LIGHT_RED,
    )

    doc.add_heading("8.1 Identity、Platt 與 Beta 的定位", level=2)
    calibration_rows = [
        ("Identity", "g(p)=p；原始機率控制組", "目前保留；production_model_changed=false。"),
        ("Monotone Platt", "sigmoid(a·logit(p)+b)，a≥0", "研究 challenger；逐 issue date strict walk-forward fit。"),
        ("Monotone Beta", "sigmoid(α·log(p)-β·log(1-p)+c)，α,β≥0", "研究 challenger；逐 issue date strict walk-forward fit。"),
        ("訓練成熟規則", "target_date < current_issue_date；同日 batch 共用同一 fit", "避免把當天才成熟的 outcome 洩漏進當天校準器。"),
        ("隔離規則", "每個 model_version × horizon 分開；同組可跨 symbol pooling", "不可跨 horizon 混 fit；pooling 假設仍需敏感度分析。"),
        ("目前決策", "六組 gate 均 keep_identity", "Platt/Beta 的 Brier／log loss 點估計改善，不足以抵銷 AUC／Recall 惡化與 CI／vintage blockers。"),
    ]
    _add_table(
        doc,
        ("方法／規則", "定義", "目前狀態"),
        calibration_rows,
        (3.2, 5.6, 6.2),
        font_size=7.3,
    )

    doc.add_heading("九、發布判定與下一步", level=1)
    _add_note(
        doc,
        "現在的發布決策",
        "NO-GO：維持 identity、不得直接 production。先建立 exact-vintage replay 與 matured live/shadow evidence，"
        "再以每個 horizon 的 BSS／log-loss CI、AUC、AP、BA、MCC、coverage、成本與風險共同審查。",
        fill=LIGHT_RED,
    )
    _add_bullet(doc, "第一優先：建立不可回寫的 forecast ledger、exact historical vintage 與完整 forecast-time baseline。")
    _add_bullet(doc, "第二優先：建立 feature-based challenger，使用 chronological walk-forward／purged evaluation，所有 threshold 只在 training fold 決定。")
    _add_bullet(doc, "第三優先：達到 BSS >0 且 paired CI 下界 >0，再追求 BSS ≥1%、AUC ≥0.53 與 AP relative lift ≥5%。")
    _add_bullet(doc, "第四優先：先 shadow，不直接上線；納入交易成本、滑價、最大回撤、錯誤代價與人工覆核。")

    doc.add_heading("十、方法依據與重現資訊", level=1)
    method_rows = [
        ("Brier / proper scores", "Brier (1950); Gneiting & Raftery (2007)", "模型選擇優先看 proper score 與 forecast-time baseline skill"),
        ("ROC-AUC", "Fawcett (2006)", "和 0.5 及其 clustered CI 比較"),
        ("Precision-Recall / AP", "Davis & Goadrich (2006)", "AP 必須和同切片 prevalence 一起看"),
        ("SHAP", "Lundberg & Lee (2017)", "只作 feature attribution，不取代 OOS 效能驗證"),
        ("Replay artifact", "重跑後輸出至 reports/forecast-model-metrics-replay.json（未隨 repo 提交）", "Machine-readable scorecard；report schema／serializer 變更會改變 file hash"),
        ("Input manifest", "9315c60b65772ed5ae495a5442ead39df06f14b5b329e7bc8d77baaf7338d000", "本次 15 個來源 CSV 的 aggregate SHA-256 manifest"),
    ]
    _add_table(
        doc,
        ("主題", "依據／Artifact", "本專案使用方式"),
        method_rows,
        (2.7, 5.0, 7.3),
        font_size=6.9,
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(footer.add_run(f"crypto-quant｜預測模型指標與參考門檻｜{REPORT_DATE}"), size=8, color="777777")
    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT)
    print(f"寫出: {OUT} ({OUT.stat().st_size / 1024:.0f} KB)")


if __name__ == "__main__":
    main()
