# -*- coding: utf-8 -*-
"""
md2docx.py — Markdown → Word(.docx) 轉換引擎（本機無 pandoc 用）。

產生 docs/ 下兩份交接文件（來源在 docs/src/）：
  docs/src/crypto-quant_系統規格書.md → docs/crypto-quant_系統規格書.docx
  docs/src/crypto-quant_交接手冊.md   → docs/crypto-quant_交接手冊.docx
一鍵重生請跑 scripts/build_docs.py（會再用 Word 更新目錄頁碼）。

排版主題：深藍標題階層＋底線、表格深藍表頭白字＋隔行網底、程式碼盒、引言左側邊條、
每個「# 章」從新頁開始（首個 # 為文件標題）、頁尾「第 X 頁／共 Y 頁」、A4、2cm 邊界。
樣式常數集中在 THEME，要換配色改那裡即可。

支援的 Markdown：
  # 標題（首個 # 為文件標題；其後每個 # 另起新頁）、## ～ ####
  段落（相鄰行以換行連接，一行就是一段最穩）
  **粗體**、*斜體*、`行內碼`、[連結](url)、<br>
  - 無序清單（縮排兩格＝下一層）、1. 有序清單、- [x] 勾選清單
  | 表格 |（首列表頭；含 |---| 分隔列）
  > 引言（連續 > 行＝多段引言）
  ``` 程式碼盒；```mermaid 以本機 mermaid-cli（npx）渲染成 PNG 嵌入，渲染失敗退回文字註記
  ![說明](相對路徑) 本地圖片；圖或圖表之後緊接以「圖：」開頭的一行＝置中圖說
  ---  分隔線
  [[TOC]] 插入 Word 目錄欄位（章節＋頁碼；build_docs.py 會用 Word 更新頁碼）
  [[PAGEBREAK]] 強制分頁

需求：python-docx；node + npx（渲染 mermaid 用；首次自動下載 @mermaid-js/mermaid-cli）。

用法（專案根目錄）：
  .venv\\Scripts\\python.exe scripts\\md2docx.py            # 轉換下方 JOBS 全部
其他腳本可： from md2docx import convert; convert(md_path, out_path)
"""
import io
import os
import re
import shutil
import struct
import subprocess
import sys
import tempfile
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")  # Windows cp950 主控台印中文用；不替換物件，避免重複包裝關閉串流
ROOT = Path(__file__).resolve().parent.parent

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

ZH_FONT = "微軟正黑體"
LATIN_FONT = "Microsoft JhengHei"
MONO = "Consolas"

# ── 配色主題 ────────────────────────────────────────────────────────────────
THEME = {
    "accent":       "1F4E79",  # 標題與表頭：深藍
    "accent_soft":  "2E6DA4",  # 次級標題
    "text":         "333333",
    "muted":        "595959",
    "table_border": "C9D6E4",
    "table_band":   "F2F6FA",  # 表格隔行網底
    "code_bg":      "F7F8FA",
    "code_border":  "E1E4E8",
    "code_ink":     "EFF1F4",  # 行內 code 底色
    "quote_bg":     "F7FAFC",
    "rule":         "D0D7DE",
    "link":         "0563C1",
    "h4":           "2D3748",
}

JOBS = [
    (ROOT / "docs" / "src" / "crypto-quant_系統規格書.md", ROOT / "docs" / "crypto-quant_系統規格書.docx"),
    (ROOT / "docs" / "src" / "crypto-quant_交接手冊.md",   ROOT / "docs" / "crypto-quant_交接手冊.docx"),
]


def _rgb(key):
    return RGBColor.from_string(THEME[key])


# ── 低階 XML 小工具 ──────────────────────────────────────────────────────────
def _set_ea(run, latin=None):
    run.font.name = latin or LATIN_FONT
    rPr = run._element.get_or_add_rPr()
    rf = rPr.rFonts if rPr.rFonts is not None else rPr._add_rFonts()
    rf.set(qn("w:eastAsia"), ZH_FONT)


def _run_shd(run, fill):
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    rPr.append(shd)


def _para_shd(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _para_border(p, edges):
    """edges = {"bottom": (size, color, space), "left": (...)}；size 為 1/8 pt。"""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge, (size, color, space) in edges.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), str(space))
        pBdr.append(el)
    pPr.append(pBdr)


def _cell_shd(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _cell_margins(cell, top=70, bottom=70, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, v in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(v))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _table_theme(t):
    """全寬、細框線。"""
    tblPr = t._tbl.tblPr
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), "5000")
    w.set(qn("w:type"), "pct")
    tblPr.append(w)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), THEME["table_border"])
        borders.append(el)
    tblPr.append(borders)


def _repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _add_field(p, instr, size=8, placeholder=None):
    """欄位（PAGE / NUMPAGES / TOC）。placeholder 為欄位尚未更新前顯示的文字。"""
    def _run(size_pt):
        r = p.add_run()
        _set_ea(r)
        r.font.size = Pt(size_pt)
        r.font.color.rgb = _rgb("muted")
        return r
    r = _run(size)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    if placeholder is not None:
        fld.set(qn("w:dirty"), "true")
    r._r.append(fld)
    r = _run(size)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = f" {instr} "
    r._r.append(it)
    r = _run(size)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    r._r.append(fld)
    if placeholder is not None:
        r = _run(size)
        r.text = placeholder
    r = _run(size)
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    r._r.append(fld)


# ── 文件骨架 ────────────────────────────────────────────────────────────────
def _new_doc():
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = s.bottom_margin = Cm(2)
        s.left_margin = s.right_margin = Cm(2.2)
    st = doc.styles["Normal"]
    st.font.name = LATIN_FONT
    st.font.size = Pt(10.5)
    st.font.color.rgb = _rgb("text")
    st.element.rPr.rFonts.set(qn("w:eastAsia"), ZH_FONT)
    st.paragraph_format.space_after = Pt(6)
    # 標題階層：16 深藍 / 14 深藍 / 13 次藍 / 11.5 深灰
    for h, sz, colkey, before, after in (
        ("Heading 1", 16, "accent", 12, 10),
        ("Heading 2", 14, "accent", 14, 7),
        ("Heading 3", 13, "accent_soft", 13, 6),
        ("Heading 4", 11.5, "h4", 10, 5),
    ):
        try:
            s = doc.styles[h]
            s.font.name = LATIN_FONT
            s.font.size = Pt(sz)
            s.font.bold = True
            s.font.color.rgb = _rgb(colkey)
            if s.element.rPr is None:
                s.element.get_or_add_rPr()
            rf = (s.element.rPr.rFonts if s.element.rPr.rFonts is not None
                  else s.element.rPr._add_rFonts())
            rf.set(qn("w:eastAsia"), ZH_FONT)
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(after)
            s.paragraph_format.keep_with_next = True
        except KeyError:
            pass
    # 目錄標題：外觀同 Heading 1，但不是大綱層級（不會把自己列進目錄）
    toc_h = doc.styles.add_style("TOC Heading Custom", WD_STYLE_TYPE.PARAGRAPH)
    toc_h.base_style = doc.styles["Normal"]
    toc_h.font.name = LATIN_FONT
    toc_h.font.size = Pt(16)
    toc_h.font.bold = True
    toc_h.font.color.rgb = _rgb("accent")
    toc_h.element.get_or_add_rPr()
    rf = (toc_h.element.rPr.rFonts if toc_h.element.rPr.rFonts is not None
          else toc_h.element.rPr._add_rFonts())
    rf.set(qn("w:eastAsia"), ZH_FONT)
    toc_h.paragraph_format.space_before = Pt(12)
    toc_h.paragraph_format.space_after = Pt(10)
    toc_h.paragraph_format.keep_with_next = True
    # 目錄各層樣式（Word 更新目錄時套用）
    for name, indent, size in (("TOC 1", 0, 10.5), ("TOC 2", 0.6, 10), ("TOC 3", 1.2, 9.5)):
        try:
            s = doc.styles[name]
        except KeyError:
            s = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = doc.styles["Normal"]
        s.font.name = LATIN_FONT
        s.font.size = Pt(size)
        s.font.bold = (name == "TOC 1")
        s.font.color.rgb = _rgb("text")
        s.element.get_or_add_rPr()
        rf = (s.element.rPr.rFonts if s.element.rPr.rFonts is not None
              else s.element.rPr._add_rFonts())
        rf.set(qn("w:eastAsia"), ZH_FONT)
        s.paragraph_format.left_indent = Cm(indent)
        s.paragraph_format.space_after = Pt(2)
    return doc


def add_footer(doc, title):
    """頁尾：細頂線＋「標題 · 第 X 頁／共 Y 頁」置中。"""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_border(p, {"top": (4, THEME["rule"], 4)})
    for txt, fld in ((f"{title}　·　第 ", "PAGE"), (" 頁／共 ", "NUMPAGES"), (" 頁", None)):
        r = p.add_run(txt)
        _set_ea(r)
        r.font.size = Pt(8)
        r.font.color.rgb = _rgb("muted")
        if fld:
            _add_field(p, fld)


def mark_update_fields(docx_path: Path):
    """在 settings.xml 加 updateFields，讓 Word 開啟時自動更新目錄（本機無 Word 時的退路）。"""
    doc = Document(str(docx_path))
    settings = doc.settings.element
    for el in settings.findall(qn("w:updateFields")):
        settings.remove(el)
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)
    doc.save(str(docx_path))


# ── 行內語法：**粗體**、*斜體*、`code`、[連結](url)、<br> ─────────────────────
_INLINE_RE = re.compile(
    r"(`[^`]+`)"                      # 1 行內碼
    r"|(\*\*[^*]+?\*\*)"              # 2 粗體
    r"|(\*(?!\s)[^*\n]+?(?<!\s)\*)"   # 3 斜體
    r"|(\[[^\]]+\]\([^)]+\))"         # 4 連結
    r"|(<br\s*/?>)"                   # 5 換行
)
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _plain_run(p, text, color=None, base_size=None, bold=False, italic=False):
    if not text:
        return
    r = p.add_run(text)
    _set_ea(r)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if color is not None:
        r.font.color.rgb = color
    if base_size is not None:
        r.font.size = Pt(base_size)
    return r


def add_inline(p, text, color=None, base_size=None, bold=False, italic=False):
    """把一段含行內語法的文字寫進段落 p。"""
    text = re.sub(r"</?(?:b|strong|i|em|u|span|sub|sup|code|kbd)[^>]*>", "", text)
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _plain_run(p, text[pos:m.start()], color, base_size, bold, italic)
        tok = m.group(0)
        if m.group(1):
            r = p.add_run(tok[1:-1])
            _set_ea(r, MONO)
            r.font.size = Pt((base_size or 10.5) - 1)
            if bold:
                r.bold = True
            if color is None:  # 深色表頭上不上灰底
                r.font.color.rgb = RGBColor.from_string("24292F")
                _run_shd(r, THEME["code_ink"])
            else:
                r.font.color.rgb = color
        elif m.group(2):
            add_inline(p, tok[2:-2], color, base_size, True, italic)
        elif m.group(3):
            add_inline(p, tok[1:-1], color, base_size, bold, True)
        elif m.group(4):
            lm = _LINK_RE.match(tok)
            r = p.add_run(re.sub(r"[*`]", "", lm.group(1)))
            _set_ea(r)
            r.font.color.rgb = _rgb("link") if color is None else color
            r.underline = True
            if base_size is not None:
                r.font.size = Pt(base_size)
        elif m.group(5):
            p.add_run().add_break()
        pos = m.end()
    if pos < len(text):
        _plain_run(p, text[pos:], color, base_size, bold, italic)


# ── 區塊 ────────────────────────────────────────────────────────────────────
def _cells(line):
    parts = line.strip().split("|")
    if parts and parts[0].strip() == "":
        parts = parts[1:]
    if parts and parts[-1].strip() == "":
        parts = parts[:-1]
    return [c.strip() for c in parts]


def _is_sep(line):
    return bool(re.match(r"^\s*\|?\s*:?-{2,}.*$", line)) and set(line.strip()) <= set("|-: ")


def _aligns(sep_line, ncols):
    out = []
    for c in _cells(sep_line):
        c = c.strip()
        if c.startswith(":") and c.endswith(":"):
            out.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif c.endswith(":"):
            out.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            out.append(WD_ALIGN_PARAGRAPH.LEFT)
    while len(out) < ncols:
        out.append(WD_ALIGN_PARAGRAPH.LEFT)
    return out


def add_table(doc, rows, aligns=None):
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    _table_theme(t)
    aligns = aligns or [WD_ALIGN_PARAGRAPH.LEFT] * ncols
    for ri, row in enumerate(rows):
        if ri == 0:
            _repeat_header(t.rows[0])
        for ci in range(ncols):
            cell = t.cell(ri, ci)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _cell_margins(cell)
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.alignment = aligns[ci]
            txt = row[ci] if ci < len(row) else ""
            if ri == 0:
                _cell_shd(cell, THEME["accent"])
                add_inline(para, txt, color=RGBColor.from_string("FFFFFF"), base_size=9.5, bold=True)
            else:
                if ri % 2 == 0:  # 隔行網底
                    _cell_shd(cell, THEME["table_band"])
                add_inline(para, txt, base_size=9)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def add_code(doc, lines):
    """程式碼盒：單格表格、淺底細框、等寬 9pt。"""
    t = doc.add_table(rows=1, cols=1)
    tblPr = t._tbl.tblPr
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), "5000")
    w.set(qn("w:type"), "pct")
    tblPr.append(w)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), THEME["code_border"])
        borders.append(el)
    tblPr.append(borders)
    cell = t.cell(0, 0)
    _cell_shd(cell, THEME["code_bg"])
    _cell_margins(cell, top=110, bottom=110, left=150, right=150)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for k, ln in enumerate(lines):
        if k:
            p.add_run().add_break()
        r = p.add_run(ln if ln else " ")
        _set_ea(r, MONO)
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("24292F")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def add_quote(doc, text):
    """引言：左側深藍邊條＋淺底＋縮排。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(3)
    _para_border(p, {"left": (18, THEME["accent"], 8)})
    _para_shd(p, THEME["quote_bg"])
    add_inline(p, text, color=RGBColor.from_string("404040"))
    return p


def add_list_item(doc, text, level=0, marker="•", ordered=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7 + 0.6 * level)
    p.paragraph_format.first_line_indent = Cm(-0.45 if ordered else -0.35)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{marker}  " if not ordered else f"{marker} ")
    _set_ea(r)
    r.font.color.rgb = _rgb("accent")
    r.bold = True
    add_inline(p, text)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    _para_border(p, {"bottom": (6, THEME["rule"], 1)})


def add_caption(doc, text):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(10)
    add_inline(cap, text, color=_rgb("muted"), base_size=9, italic=True)


def add_diagram_note(doc, label="流程圖"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f"〔{label}：本機未能渲染圖表，請於 GitHub 線上版檢視〕")
    _set_ea(r)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string("888888")


def add_toc(doc, title="目錄（章節與頁碼）", levels="1-2"):
    h = doc.add_paragraph(style="TOC Heading Custom")
    add_inline(h, title, color=_rgb("accent"), base_size=16, bold=True)
    p = doc.add_paragraph()
    _add_field(p, f'TOC \\o "{levels}" \\h \\z \\u', size=10,
               placeholder="（目錄需在 Word 中更新：按 Ctrl+A 再按 F9）")


# ── Mermaid → PNG（本機 mermaid-cli via npx；失敗回 None）─────────────────────
_MMDC_DEAD = False
_MERMAID_CONFIG = (
    '{"theme":"neutral","themeVariables":{"fontSize":"18px",'
    '"fontFamily":"Microsoft JhengHei, Noto Sans CJK TC, sans-serif"},'
    '"flowchart":{"htmlLabels":true,"curve":"basis","nodeSpacing":40,"rankSpacing":50}}'
)


def _png_size(path):
    try:
        with open(path, "rb") as f:
            head = f.read(24)
        if head[:8] == b"\x89PNG\r\n\x1a\n":
            return struct.unpack(">II", head[16:24])
    except Exception:
        pass
    return None


def _render_mermaid(src, idx, tmp: Path):
    global _MMDC_DEAD
    if _MMDC_DEAD:
        return None
    mmd = tmp / f"d{idx}.mmd"
    png = tmp / f"d{idx}.png"
    mmd.write_text(src, encoding="utf-8")
    pconf = tmp / "pconf.json"
    if not pconf.exists():
        pconf.write_text('{"args":["--no-sandbox"]}', encoding="utf-8")
    mconf = tmp / "mconf.json"
    if not mconf.exists():
        mconf.write_text(_MERMAID_CONFIG, encoding="utf-8")
    cmd = (f'npx -y @mermaid-js/mermaid-cli -i "{mmd}" -o "{png}" '
           f'-b white -s 3 -w 1400 -p "{pconf}" -c "{mconf}"')
    try:
        subprocess.run(cmd, shell=True, capture_output=True, timeout=240)
    except Exception:
        _MMDC_DEAD = True
        return None
    if not png.exists():
        if idx == 1:
            _MMDC_DEAD = True
        return None
    return png


def _add_centered_picture(doc, img_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    sz = _png_size(img_path)
    if sz and sz[1] > sz[0] * 1.35:
        run.add_picture(str(img_path), height=Cm(17))
    else:
        run.add_picture(str(img_path), width=Cm(15.8))


def add_diagram(doc, src, idx, tmp):
    png = _render_mermaid(src, idx, tmp)
    if not png:
        add_diagram_note(doc)
        return False
    _add_centered_picture(doc, png)
    return True


# ── 主轉換 ──────────────────────────────────────────────────────────────────
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*#*\s*$")
_ULIST_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
_OLIST_RE = re.compile(r"^(\s*)(\d+)[.)]\s+(.*)$")
_TASK_RE = re.compile(r"^\[( |x|X)\]\s+(.*)$")
_IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def convert(src: Path, out: Path):
    """把單一 .md 轉成 .docx（含 mermaid 渲染）。回傳 (mermaid總數, 成功渲染數, 是否含目錄欄位)。"""
    md = Path(src).read_text(encoding="utf-8").splitlines()
    doc = _new_doc()
    tmp = Path(tempfile.mkdtemp(prefix="mmd_"))
    mermaid_n = rendered = 0
    has_toc = False
    h1_seen = 0
    doc_title = Path(src).stem
    i, n = 0, len(md)
    last_was_figure = False  # 用來辨識緊接在圖後面的「圖：」圖說
    try:
        while i < n:
            line = md[i]
            stripped = line.strip()

            # 程式碼區塊
            if stripped.startswith("```"):
                lang = stripped[3:].strip().lower()
                buf = []
                i += 1
                while i < n and not md[i].strip().startswith("```"):
                    buf.append(md[i])
                    i += 1
                i += 1  # 跳過結尾 ```
                if lang == "mermaid":
                    mermaid_n += 1
                    if add_diagram(doc, "\n".join(buf), mermaid_n, tmp):
                        rendered += 1
                    last_was_figure = True
                else:
                    add_code(doc, buf)
                    last_was_figure = False
                continue

            # 空行
            if not stripped:
                i += 1
                continue

            # 特殊標記
            if stripped == "[[TOC]]":
                add_toc(doc)
                has_toc = True
                i += 1
                continue
            if stripped == "[[PAGEBREAK]]":
                doc.add_page_break()
                i += 1
                continue

            # 表格
            if stripped.startswith("|") and "|" in stripped[1:]:
                block, aligns = [], None
                while i < n and md[i].strip().startswith("|"):
                    if _is_sep(md[i]):
                        aligns = md[i]
                    else:
                        block.append(_cells(md[i]))
                    i += 1
                if block:
                    add_table(doc, block, _aligns(aligns, max(len(r) for r in block)) if aligns else None)
                last_was_figure = False
                continue

            # 圖片
            mi = _IMG_RE.match(stripped)
            if mi:
                img_path = (Path(src).resolve().parent / mi.group(2)).resolve()
                if img_path.exists():
                    _add_centered_picture(doc, img_path)
                    if mi.group(1):
                        add_caption(doc, mi.group(1))
                        last_was_figure = False
                    else:
                        last_was_figure = True
                else:
                    add_diagram_note(doc, label=f"圖：{mi.group(1) or mi.group(2)}（本機檔案未找到）")
                    last_was_figure = False
                i += 1
                continue

            # 標題（setext：下一行為 === 也算文件標題）
            m = _HEADING_RE.match(line)
            setext = (not m and i + 1 < n and re.match(r"^=+\s*$", md[i + 1]) and h1_seen == 0)
            if m or setext:
                if setext:
                    depth, text = 1, stripped
                    i += 1
                else:
                    depth, text = len(m.group(1)), m.group(2)
                if depth == 1:
                    h1_seen += 1
                    if h1_seen == 1:
                        doc_title = re.sub(r"[*`]", "", text).strip()
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(12)
                        _para_border(p, {"bottom": (12, THEME["accent"], 6)})
                        add_inline(p, text, color=_rgb("accent"), base_size=22, bold=True)
                    else:
                        h = doc.add_heading(level=1)
                        h.paragraph_format.page_break_before = True
                        _para_border(h, {"bottom": (8, THEME["accent"], 4)})
                        add_inline(h, text, color=_rgb("accent"))
                else:
                    h = doc.add_heading(level=min(depth, 4))
                    add_inline(h, text)
                last_was_figure = False
                i += 1
                continue

            # 分隔線
            if re.match(r"^(-{3,}|\*{3,}|_{3,})\s*$", stripped):
                add_hr(doc)
                i += 1
                continue

            # 引言（連續 > 行；空的 > 行分段）
            if stripped.startswith(">"):
                buf = []
                while i < n and md[i].strip().startswith(">"):
                    body = md[i].strip()[1:].strip()
                    if body:
                        buf.append(body)
                    elif buf and buf[-1] != "":
                        buf.append("")
                    i += 1
                para, quotes = [], []
                for b in buf:
                    if b == "":
                        if para:
                            quotes.append("<br>".join(para))
                            para = []
                    else:
                        para.append(b)
                if para:
                    quotes.append("<br>".join(para))
                for q in quotes:
                    add_quote(doc, q)
                last_was_figure = False
                continue

            # 清單（無序／有序／勾選；縮排兩格＝下一層）
            mu, mo = _ULIST_RE.match(line), _OLIST_RE.match(line)
            if mu or mo:
                while i < n:
                    ln = md[i]
                    mu, mo = _ULIST_RE.match(ln), _OLIST_RE.match(ln)
                    if not (mu or mo):
                        break
                    if mu:
                        level = len(mu.group(1).replace("\t", "  ")) // 2
                        text = mu.group(2)
                        mt = _TASK_RE.match(text)
                        if mt:
                            add_list_item(doc, mt.group(2), level, "☑" if mt.group(1).lower() == "x" else "☐")
                        else:
                            add_list_item(doc, text, level, "•" if level == 0 else "◦")
                    else:
                        level = len(mo.group(1).replace("\t", "  ")) // 2
                        add_list_item(doc, mo.group(3), level, f"{mo.group(2)}.", ordered=True)
                    i += 1
                last_was_figure = False
                continue

            # 圖說（緊接在圖之後、以「圖：」開頭）
            if last_was_figure and (stripped.startswith("圖：") or stripped.startswith("圖 ")):
                add_caption(doc, stripped)
                last_was_figure = False
                i += 1
                continue

            # 一般段落：相鄰行合併為一段（以換行連接）
            buf = [stripped]
            i += 1
            while i < n:
                nxt = md[i]
                ns = nxt.strip()
                if (not ns or ns.startswith(("```", "|", ">", "#", "!["))
                        or _ULIST_RE.match(nxt) or _OLIST_RE.match(nxt)
                        or re.match(r"^(-{3,}|\*{3,}|_{3,})\s*$", ns)
                        or ns in ("[[TOC]]", "[[PAGEBREAK]]")):
                    break
                buf.append(ns)
                i += 1
            add_inline(doc.add_paragraph(), "<br>".join(buf))
            last_was_figure = False
        add_footer(doc, doc_title)
        doc.core_properties.title = doc_title
        doc.core_properties.author = "crypto-quant"
        doc.save(str(out))
    finally:
        shutil.rmtree(tmp, ignore_errors=True)
    return mermaid_n, rendered, has_toc


def main():
    for src, out in JOBS:
        if not src.exists():
            print(f"跳過（來源不存在）：{src}")
            continue
        mn, rn, toc = convert(src, out)
        tag = f"（{rn}/{mn} 圖）" if mn else ""
        if mn and rn == 0:
            tag += " ⚠️ 渲染失敗，退回文字註記"
        if toc:
            tag += "（含目錄欄位，需 build_docs.py 更新頁碼）"
        print(f"寫出: {out.name}  ({os.path.getsize(out)/1024:.0f} KB) {tag}")


if __name__ == "__main__":
    main()
